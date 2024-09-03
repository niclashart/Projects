import docx
from docx.shared import Pt 
from datetime import datetime
import pypandoc
import requests

def hole_alle_patienten(token):
    url = "https://api.medico-system.com/patienten"
    headers = {
        'Authorization': f'Bearer {token}',
        'Content-Type': 'application/json'
    }
    response = requests.get(url, headers=headers)
    
    if response.status_code == 200:
        return response.json()  # Annahme: Die API gibt eine Liste von Patienten zurück
    else:
        raise Exception(f"Fehler beim Abruf der Patientendaten: {response.status_code}")

def erstelle_arztbrief(patient_data):
    # Patientendaten
    patient_name = patient_data['name']
    geburtsdatum = patient_data['geburtsdatum']
    anamnese = patient_data['anamnese']
    diagnosen = patient_data['diagnosen']
    untersuchungen = patient_data['untersuchungen']
    therapien = patient_data['therapien']
    medikation = patient_data['medikation']
    
    # Erstelle Word Dokument
    doc = docx.Document()
    doc.add_heading('Arztbrief', level=1)
    datum = datetime.today().strftime('%d.%m.%Y')
    doc.add_paragraph(f"Name des Patienten: {patient_name}")
    doc.add_paragraph(f"Geburtsdatum: {geburtsdatum}")
    doc.add_paragraph(f"Datum: {datum}")
    doc.add_heading('Anamnese:', level=2)
    doc.add_paragraph(anamnese)
    doc.add_heading('Diagnosen:', level=2)
    for diagnose in diagnosen:
        doc.add_paragraph(diagnose, style='List Bullet')
    doc.add_heading('Untersuchungsergebnisse:', level=2)
    doc.add_paragraph(untersuchungen)
    doc.add_heading('Therapieempfehlungen:', level=2)
    doc.add_paragraph(therapien)
    doc.add_paragraph("\nMit freundlichen Grüßen,\n\n[Arzt Name]\n[Klinik/Institution]")

    # Dateinamen generieren
    word_datei = f"arztbrief_{patient_name.replace(' ', '_')}.docx"
    doc.save(word_datei)
    return word_datei

def konvertiere_word_zu_pdf(word_datei):
    pdf_datei = word_datei.replace('.docx', '.pdf')
    output = pypandoc.convert_file(word_datei, '.pdf', outputfile=pdf_datei)
    assert output == '', 'Fehler bei der Konvertierung'
    return pdf_datei

def erstelle_arztbriefe_fuer_alle_patienten():
    # Authentifizierungstoken (dieser muss vorher vom KIS-System bereitgestellt werden)
    token = "DEIN_AUTH_TOKEN"
    
    # Abrufen aller Patienten
    patienten = hole_alle_patienten(token)
    
    # Schleife über alle Patienten und Arztbrief erstellen
    for patient in patienten:
        try:
            # Word-Dokument erstellen
            word_datei = erstelle_arztbrief(patient)
            print(f"Word-Dokument für {patient['name']} erstellt: {word_datei}")
            
            # Konvertierung in PDF
            pdf_datei = konvertiere_word_zu_pdf(word_datei)
            print(f"PDF-Dokument für {patient['name']} erstellt: {pdf_datei}")
        
        except Exception as e:
            print(f"Fehler bei der Erstellung des Arztbriefes für {patient['name']}: {e}")

# Skript starten
erstelle_arztbriefe_fuer_alle_patienten()

